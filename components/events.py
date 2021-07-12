from docx import Document
import os
import time
import tkinter as tk
from tkinter import filedialog


l = []


def clearlist(): # funtion to clean the KO list
    list = l.clear() # cleaning the list
    return list # returnin the list

def removedup(lista:list): # funtion to remove duplicated values of the list
    l = []
    for i in lista: # loop to iterate the list values
        if i not in l: # assign values to list if the value is not in it
            l.append(i) # append the value to the list
    return sorted(l) # sorting and returning the value


def ko(ct:int, flag:str): # function to create a list of KOs - 2 values assigned, 1st the number of spec, 2nd a flag to determine the function    
    if flag == '0': # comparing the flag value
        l.append(ct) # appeding the value "ct" to list
        list = removedup(l) # calling the funtion to remove a possible duplicated value
    elif flag == '1': # comparing the flag value
        list = clearlist() # calling the funtion to clean the list
    return list # returning the list 


def removevalue(ct:int): # funtion to remove the KO of the list
    l.remove(ct) # removing the spec of the list
    list = removedup(l) # calling the funtion to remove a possible duplicated value
    return list # returning the list



def remove_row(table, row): # funtion to remove rows from tables (lines in the file)
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr) # remove lines from file
   
   
def def_path(): # funtion to choose a file 
    root = tk.Tk() 
    root.withdraw()
    file_path = filedialog.askopenfilename(initialdir=os.getcwd(), title="Select file", filetypes=[("CFTS Files", "*.docm")])
    if file_path.endswith(".docm"):
        return os.path.abspath(file_path)
    else:
        return 'All .docm files in root folder will be filtered'
   
    
d = {}
def create_dict(c,cell): # function to create a dict 
    c = str(c) # converting the value to string
    d[c] = {'title':cell, 'sts':"OK"} # assing the spec and status to the dictionary
    return d # returning the dictionary


#filtering and saving document
def filtering(arch, variant, ECU, file_path):
    file = (os.path.basename(file_path))
    d.clear()
#if __name__ == "__main__":
    #for file in os.listdir(os.getcwd()):
    if file.endswith(".docm"):
        document = Document(file)
        e=1        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for index, paragraph in enumerate(cell.paragraphs):
                        if index == 0:                        
                            if ':]' in paragraph.text:
                                k=0
                                for c in paragraph.text:
                                    if ']' in c:
                                        k=k+1
                                if k == 1:                            
                                    if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) or variant in paragraph.text or ECU in paragraph.text or 'LATAM' in paragraph.text: 
                                        dict1 = create_dict(e,cell.text) 
                                        e += 1
                                    else:
                                        remove_row(table, row)
                                if k == 2:                            
                                    if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and ECU in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and 'LATAM' in paragraph.text or variant in paragraph.text and ECU in paragraph.text or variant in paragraph.text and 'LATAM' in paragraph.text or ECU in paragraph.text and 'LATAM' in paragraph.text:
                                        dict1 = create_dict(e,cell.text) 
                                        e += 1
                                    else:
                                        #print('2- delete - '+paragraph.text)
                                        remove_row(table, row)
                                if k == 3:                            
                                    if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text and ECU in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text and 'LATAM' in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and ECU in paragraph.text and 'LATAM' in paragraph.text or variant in paragraph.text and ECU in paragraph.text  and 'LATAM' in paragraph.text:
                                        dict1 = create_dict(e,cell.text) 
                                        e += 1
                                    else:
                                        #print('3- delete - '+paragraph.text)
                                        remove_row(table, row)
                                if k == 4:                            
                                    if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text and ECU in paragraph.text and 'LATAM' in paragraph.text :
                                        dict1 = create_dict(e,cell.text) 
                                        e += 1
                                    else:
                                        #print('4- delete - '+paragraph.text)
                                        remove_row(table, row)
        document.save('FILTRADO_'+file) # saving the filtered file 
        return dict1 # returnig the dictionary
 
    
#filtering and saving all documents of the folder
def filtering_orig(arch, variant, ECU):
    for file in os.listdir(os.getcwd()):
        if file.endswith(".docm"):
            document = Document(file)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for index, paragraph in enumerate(cell.paragraphs):
                            if index == 0:                        
                                if ':]' in paragraph.text:
                                    k=0
                                    for c in paragraph.text:
                                        if ']' in c:
                                            k=k+1
                                    if k == 1:                            
                                        if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) or variant in paragraph.text or ECU in paragraph.text or 'LATAM' in paragraph.text:
                                            pass
                                            #print('1- '+paragraph.text)                                    
                                        else:
                                            #print('1- delete - '+paragraph.text)
                                            remove_row(table, row)
                                    if k == 2:                            
                                        if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and ECU in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and 'LATAM' in paragraph.text or variant in paragraph.text and ECU in paragraph.text or variant in paragraph.text and 'LATAM' in paragraph.text or ECU in paragraph.text and 'LATAM' in paragraph.text:
                                            pass
                                            #print('2 - '+paragraph.text)
                                        else:
                                            #print('2- delete - '+paragraph.text)
                                            remove_row(table, row)
                                    if k == 3:                            
                                        if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text and ECU in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text and 'LATAM' in paragraph.text or ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and ECU in paragraph.text and 'LATAM' in paragraph.text or variant in paragraph.text and ECU in paragraph.text  and 'LATAM' in paragraph.text:
                                            pass
                                            #print('3 - '+paragraph.text)
                                        else:
                                            #print('3- delete - '+paragraph.text)
                                            remove_row(table, row)
                                    if k == 4:                            
                                        if ('allSys: CTS1_2' in paragraph.text or arch in paragraph.text) and variant in paragraph.text and ECU in paragraph.text and 'LATAM' in paragraph.text :
                                            pass
                                            #print('4 - '+paragraph.text)
                                        else:
                                            #print('4- delete - '+paragraph.text)
                                            remove_row(table, row)
            document.save('FILTRADO_'+file)